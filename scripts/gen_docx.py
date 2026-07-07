#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
gen_docx.py — 将本技能生成的 Markdown 技术文档导出为 Word (.docx)。

用法 / Usage:
    python gen_docx.py input.md output.docx [--title "标题"]

依赖 / Dependencies:
    pip install python-docx

说明:
    - 支持标题(#~######)、段落、无序/有序列表、表格(GFM 管道表格)、代码块。
    - Mermaid 代码块会以等宽文本保留，并附提示"此为拓扑/原理图源码，可用 Mermaid 渲染"。
    - 复杂图形建议用专业 EDA (KiCad/Altium) 出正式原理图，本导出仅供文档归档。
"""
import sys
import re
import argparse


def add_table(doc, header, rows):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h.strip()
    for row in rows:
        cells = table.add_row().cells
        for i, c in enumerate(row):
            if i < len(cells):
                cells[i].text = c.strip()
    return table


def parse_table_block(lines, start):
    """Parse a GFM pipe table starting at index `start`. Returns (header, rows, next_index)."""
    def split_row(line):
        line = line.strip().strip('|')
        return [c.strip() for c in line.split('|')]

    header = split_row(lines[start])
    # separator line (---|---)
    rows = []
    i = start + 2
    while i < len(lines) and lines[i].strip().startswith('|'):
        rows.append(split_row(lines[i]))
        i += 1
    return header, rows, i


def is_table_sep(line):
    return bool(re.match(r'^\s*\|?[\s:\-|]+\|?\s*$', line)) and '-' in line


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('input')
    ap.add_argument('output')
    ap.add_argument('--title', default=None)
    args = ap.parse_args()

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("需要安装 python-docx: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    with open(args.input, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = Document()
    if args.title:
        doc.add_heading(args.title, level=0)

    i = 0
    in_code = False
    code_lang = ''
    code_buf = []

    while i < len(lines):
        line = lines[i]

        # code block fences
        m = re.match(r'^```(\w*)', line)
        if m and not in_code:
            in_code = True
            code_lang = m.group(1)
            code_buf = []
            i += 1
            continue
        if line.strip() == '```' and in_code:
            in_code = False
            note = ''
            if code_lang == 'mermaid':
                note = '【拓扑/原理图 Mermaid 源码，可用 Mermaid 渲染为图】\n'
            p = doc.add_paragraph()
            run = p.add_run(note + '\n'.join(code_buf))
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # headings
        hm = re.match(r'^(#{1,6})\s+(.*)', line)
        if hm:
            level = len(hm.group(1))
            doc.add_heading(hm.group(2).strip(), level=min(level, 4))
            i += 1
            continue

        # tables
        if line.strip().startswith('|') and i + 1 < len(lines) and is_table_sep(lines[i + 1]):
            header, rows, ni = parse_table_block(lines, i)
            add_table(doc, header, rows)
            i = ni
            continue

        # list items
        lm = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if lm:
            doc.add_paragraph(lm.group(2).strip(), style='List Bullet')
            i += 1
            continue
        om = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if om:
            doc.add_paragraph(om.group(2).strip(), style='List Number')
            i += 1
            continue

        # horizontal rule
        if re.match(r'^---+\s*$', line):
            i += 1
            continue

        # blockquote
        if line.strip().startswith('>'):
            p = doc.add_paragraph(line.strip().lstrip('>').strip())
            p.style = 'Intense Quote'
            i += 1
            continue

        # normal paragraph (strip simple markdown emphasis)
        text = line.strip()
        if text:
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            doc.add_paragraph(text)
        i += 1

    doc.save(args.output)
    print(f"已生成 / Saved: {args.output}")


if __name__ == '__main__':
    main()
